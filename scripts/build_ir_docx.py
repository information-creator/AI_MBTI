"""AI-MBTI IR/PR docx 생성 스크립트.

실행: python3 scripts/build_ir_docx.py
출력: docs/IR_AI-MBTI.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
OUT = ROOT / "docs" / "IR_AI-MBTI.docx"

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
ACCENT = RGBColor(0x2E, 0x7D, 0xFF)
GRAY = RGBColor(0x55, 0x5B, 0x66)
LIGHT = RGBColor(0xE8, 0xEE, 0xF7)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, size: int = 20, color: RGBColor = NAVY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ACCENT


def add_para(doc: Document, text: str, size: int = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bullets(doc: Document, items: list[str], size: int = 10) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0] if p.runs else p.add_run()
        run.text = it
        run.font.size = Pt(size)


def add_image(doc: Document, rel_path: str, width_cm: float = 14.0, caption: str | None = None) -> None:
    img = PUBLIC / rel_path
    if not img.exists():
        add_para(doc, f"[이미지 없음: {rel_path}]", color=GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY


def add_image_row(doc: Document, entries: list[tuple[str, str]], row_width_cm: float = 16.0) -> None:
    """entries: list of (rel_path, caption). 2~3 images side by side."""
    valid = [(PUBLIC / p, c) for p, c in entries if (PUBLIC / p).exists()]
    if not valid:
        return
    n = len(valid)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    col_w = Cm(row_width_cm / n)
    for i, (img_path, caption) in enumerate(valid):
        cell_img = tbl.cell(0, i)
        cell_img.width = col_w
        cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell_img.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(str(img_path), width=Cm(row_width_cm / n - 0.4))
        except Exception as e:
            para.add_run(f"[err {img_path.name}: {e}]").font.size = Pt(8)
        cell_cap = tbl.cell(1, i)
        cell_cap.width = col_w
        cap_p = cell_cap.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap_p.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "0B1F3A")
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Default font: 맑은 고딕
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    # ─────────────────────────────────────────────
    # 1. 표지
    # ─────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    logo = PUBLIC / "Adv" / "로고 및 참조" / "로고_transparent.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Cm(8))

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AI-MBTI")
    tr.bold = True
    tr.font.size = Pt(48)
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("AI 시대 생존력 진단 서비스")
    sr.font.size = Pt(18)
    sr.font.color.rgb = ACCENT

    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tgr = tagline.add_run("20문항으로 진단하는 16가지 AI 활용 유형 · 데이터 드리븐 마케팅")
    tgr.font.size = Pt(12)
    tgr.font.color.rgb = GRAY

    for _ in range(6):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("IR / PR Deck · mcodegc.com · 2026.04")
    mr.font.size = Pt(11)
    mr.font.color.rgb = GRAY

    page_break(doc)

    # ─────────────────────────────────────────────
    # 2. 서비스 개요
    # ─────────────────────────────────────────────
    add_heading(doc, "01. 서비스 개요", size=22)
    add_para(
        doc,
        "AI-MBTI는 사용자의 AI 활용 성향을 20문항으로 진단하여 16가지 유형 중 하나로 분류하고, "
        "유형별 맞춤 학습·부트캠프 경로를 제안하는 진단 서비스입니다. 2026년 4월 기준 mcodegc.com에서 운영 중이며, "
        "Next.js 16 · React 19 · Supabase 기반으로 구축되었습니다.",
        size=11,
    )

    add_subheading(doc, "■ 핵심 플로우")
    add_bullets(doc, [
        "/test — 20문항 응답 수집 (5차원 점수제, A~E 각 0-4점)",
        "/analyzing — 프로그레스바 + Supabase 결과 저장 병렬 실행",
        "/result/[id] — 16유형 중 하나 판정 + 부트캠프·이북·오픈챗 CTA",
    ], size=10)

    add_subheading(doc, "■ 유형 산출 로직 (5차원 → 4축 변환)")
    add_table(
        doc,
        ["차원", "축", "변환 기준"],
        [
            ["A + B 합산", "AI 활용 (A / S)", "합산 ≥ 5 → A (AI Native)"],
            ["C", "업무방식 (H / T)", "≥ 3 → H (Human-centric)"],
            ["D", "강점 (L / C)", "≥ 3 → L (Leading)"],
            ["E", "실행력 (F / P)", "≥ 3 → F (Fast)"],
        ],
    )
    add_para(doc, "최종 TypeCode 예: AHLF, ATLP, SHCP … (2 × 2 × 2 × 2 = 16유형)", size=10, color=GRAY)

    add_subheading(doc, "■ 16유형 캐릭터 시스템")
    add_image_row(doc, [
        ("characters/AI 팀장.png", "AHLF · AI 팀장"),
        ("characters/LLM 설계자.png", "AHCF · LLM 설계자"),
        ("characters/DE 전략가.png", "SHLF · DE 전략가"),
    ], row_width_cm=15)

    page_break(doc)

    # ─────────────────────────────────────────────
    # 3. 마케팅 실행 현황
    # ─────────────────────────────────────────────
    add_heading(doc, "02. 마케팅 실행 현황", size=22)
    add_para(
        doc,
        "AI-MBTI는 Meta Ads 중심의 퍼포먼스 마케팅을 전개하고 있으며, 소재별 A/B/C 테스트와 "
        "성별 타겟팅 이원화를 통해 CPL을 최적화하고 있습니다.",
        size=11,
    )

    # ─ 소재 1
    add_subheading(doc, "① Instagram 피드 광고 — 성별 타겟팅 이원화")
    add_image_row(doc, [
        ("Adv/여자/사무실 여자 완성.png", "여성 타겟 소재"),
        ("Adv/남자/사무실 남자 완성.png", "남성 타겟 소재"),
    ], row_width_cm=15)
    add_bullets(doc, [
        "직장인의 사무실 상황을 시각화하여 공감도 극대화 — 타겟의 일상에 자연스럽게 침투",
        "여성·남성 소재를 분리 제작하여 Meta Ads 오디언스별 전환율을 독립 측정",
        "동일 메시지·다른 비주얼 구조로 소재 피로도를 분산시키고 A/B 최적점을 탐색",
        "9:16(스토리) · 1:1(피드) · 1.91:1(가로) 3개 규격 동시 운영으로 전 지면 커버",
        "최근 7일 기준 CTR 3.5%대 달성 (업계 평균 1.8~2.0% 대비 약 1.8배)",
    ])

    # ─ 소재 2
    add_subheading(doc, "② 감정 소구 — Fear 캠페인")
    add_image(doc, "screen/fear/fear.png", width_cm=12, caption="AIMBTI-fear-20260422 소재")
    add_bullets(doc, [
        "“AI한테 내 일 뺏길까?” — 직장인의 AI 대체 공포를 직접적으로 자극하는 카피",
        "Simple·Social 소재 대비 클릭 유도력 검증 목적의 실험 캠페인",
        "클릭은 발생하나 Lead 전환 0건 — 공포 소구의 한계 데이터 포인트 확보",
        "48시간 내 개선 또는 중단 판단 기준 마련 (데이터 기반 소재 라이프사이클 운영)",
        "동일 예산 범위 내에서 Simple 소재로 재배분 → 전체 CPL 감소 효과",
    ])

    page_break(doc)

    # ─ 소재 3
    add_subheading(doc, "③ 랜딩페이지 버전 A/B/C 테스트")
    add_image_row(doc, [
        ("landing/v1/full.png", "랜딩 v1"),
        ("landing/v3/full.png", "랜딩 v3"),
        ("landing/v4/full.png", "랜딩 v4 (모바일 최적화)"),
    ], row_width_cm=16)
    add_bullets(doc, [
        "동일 트래픽 소스에 대해 랜딩 3버전을 병행 운영, test_start 전환율을 실시간 비교",
        "v4에서 모바일 우선 레이아웃 및 CTA 위치 개선 → 테스트 시작율 유의미 상승",
        "버전별 URL은 UTM 파라미터로 구분, GA4 + Supabase 이중 이벤트로 교차 검증",
        "소재(Ad)와 랜딩(LP)을 분리 실험하여 퍼널 각 단계의 기여도를 독립 측정",
        "주 1회 리뷰 사이클로 하위 성과 버전을 정리·대체 (지속적 이터레이션)",
    ])

    add_subheading(doc, "④ 콘텐츠 기반 오가닉 유입 (zunza 시리즈)")
    add_image_row(doc, [
        ("zunza/AISERVICE/1.png", "AI 서비스 소개"),
        ("zunza/DA/1.png", "DA 콘텐츠"),
        ("zunza/AILLM/1.png", "AI·LLM 콘텐츠"),
    ], row_width_cm=16)
    add_bullets(doc, [
        "유형별 캐릭터 IP를 활용한 인스타그램 카드뉴스 제작 — 오가닉 도달 강화",
        "진단 완료자에게 유형 캐릭터 이미지 공유 기능 제공 → 바이럴 루프 설계",
        "Meta 유료 광고와 오가닉 콘텐츠의 메시지·톤앤매너 일관성 유지",
        "콘텐츠 시리즈별 조회수·저장수 데이터를 Supabase 이벤트로 연동 관측",
        "캐릭터 IP 자산화로 추후 굿즈·캠페인 2차 활용 가능성 확보",
    ])

    page_break(doc)

    # ─────────────────────────────────────────────
    # 4. 데이터 대시보드 (핵심 강조)
    # ─────────────────────────────────────────────
    add_heading(doc, "03. 데이터 대시보드 — 의사결정의 중심", size=22, color=NAVY)
    add_para(
        doc,
        "AI-MBTI는 GA4 · Meta Ads API · Supabase를 단일 대시보드로 통합하여, "
        "모든 마케팅 의사결정을 수치에 근거하여 수행합니다. 내부 대시보드 4종을 운영 중이며, "
        "캠페인 단위부터 개별 유저 퍼널까지 drill-down 분석이 가능합니다.",
        size=11,
    )

    add_subheading(doc, "■ 통합 대시보드 v3 — 실제 운영 화면")
    add_image(doc, "screen/dashboard3.png", width_cm=16,
              caption="종합 뷰 · 퍼널 · 광고 성과 · 진단 통합 (2026-03-03 ~ 2026-04-23)")
    add_para(
        doc,
        "총 방문 4,339 · 2차 전환 390 (E2E 8.99%) · 유효 CPA ₩599 · 전자책 수강 88. "
        "Meta + Google 통합 CPC ₩57, 총 광고비 ₩233,545. "
        "진단 엔진이 퍼널 단계별 이탈률을 자동 분석하여 위험·주의·양호 라벨을 실시간 표시합니다.",
        size=10,
    )

    page_break(doc)
    add_subheading(doc, "■ 퍼널 분석 뷰 — 단계별 이탈 원인 자동 진단")
    add_image(doc, "screen/dashboard3-funnel.png", width_cm=16,
              caption="Unbounce 2024 · WordStream 2025 · LeadQuizzes 벤치마크 내장 비교")
    add_para(
        doc,
        "방문→CTA 27.6% (업계 평균 40% 대비 -31%, 3,144명 이탈) · CTA→테스트 91.0% (+1%) · "
        "테스트→완료 95.7% (+37%) · 완료→결과 93.9% (-1%) · 결과→2차 전환 39.8% (+398%). "
        "업계 평균을 내장 상수로 보관하여 실시간 자동 비교·경고.",
        size=10,
    )

    page_break(doc)
    add_subheading(doc, "■ Meta Ads 뷰 — 캠페인별 소재 성과")
    add_image(doc, "screen/dashboard3-meta.png", width_cm=16,
              caption="지출·노출·클릭·CTR·CPC·CPM·전환·CVR 8개 KPI 통합")
    add_para(
        doc,
        "활성 3개 소재 기준 지출 ₩88,564 · CTR 3.04% · CPC ₩560 · "
        "Lead 15 · CPL ₩5,904. 캠페인 단위(simple · social · fear) drill-down 가능, "
        "Meta Pixel Lead 이벤트 기반으로 소재별 실구매 전환 직접 비교.",
        size=10,
    )

    add_subheading(doc, "■ A/B 테스트 뷰 — 랜딩·소재 변형 실시간 비교")
    add_image(doc, "screen/dashboard3-abtest.png", width_cm=16,
              caption="변형별 PV · CTA · 테스트 · 2차 전환을 단일 테이블로 비교")
    add_para(
        doc,
        "A/B/C 변형을 동일 기간 내 독립 집계하여 전환율 차이를 정량 비교. "
        "Supabase events의 landing_version · ad_variant 디멘션을 기준으로 분리, "
        "유의미한 차이가 나오면 즉시 예산 재배분 판단.",
        size=10,
    )

    add_subheading(doc, "■ 대시보드 4종 구성")
    add_table(
        doc,
        ["대시보드", "경로", "핵심 지표"],
        [
            ["Overview", "/dashboard2/overview", "전체 KPI · 일별 트래픽 · Lead 추이"],
            ["Funnel", "/dashboard2/funnel", "page_view → test_start → result_view → CTA click 전환율"],
            ["Meta Ads", "/dashboard2/meta", "캠페인별 노출·클릭·CPL·Lead (Graph API 직연동)"],
            ["Google Ads", "/dashboard2/google", "GA4 Data API 기반 UTM 소스 성과"],
            ["A/B Test", "/dashboard2/ab-test", "Meta 3-캠페인 A/B/C 실시간 비교 (simple·social·fear)"],
            ["Mobile", "/dashboard4", "현장·출퇴근 중 모바일 모니터링 전용 뷰"],
        ],
    )

    add_subheading(doc, "■ 이벤트 이중 트래킹 아키텍처")
    add_para(doc, "GA4 gtag와 Supabase events 테이블에 동일 이벤트를 병렬 기록 → 신뢰도·재현성 확보.", size=10)
    add_bullets(doc, [
        "추적 이벤트: page_view · test_start · test_complete · result_view · openchat_click · ebook_click",
        "Meta Pixel Lead 이벤트는 결과 페이지 CTA 3종에서 발화 (ebook_unlock · ebook_download · openchat)",
        "Clarity 히트맵·세션리코딩으로 정성 데이터 보완 → 정량+정성 크로스체크",
    ])

    page_break(doc)

    # ─────────────────────────────────────────────
    # 5. 성과 요약 (실데이터)
    # ─────────────────────────────────────────────
    add_heading(doc, "04. 성과 요약", size=22)
    add_para(doc, "(2026-03-03 ~ 2026-04-23 · 52일 누적 · Meta Ads API 실시간)", size=10, color=GRAY)

    add_subheading(doc, "■ 캠페인별 Lead 성과 (Meta Pixel Lead · 세션당 1회 발사 · 활성 3개 캠페인)")
    add_table(
        doc,
        ["캠페인", "CTR", "CPC", "비용", "Lead", "CPL"],
        [
            ["AIMBTI-simple-v2 (최우수)", "2.96%", "₩494", "₩28,155", "7", "₩4,022"],
            ["AIMBTI-social-v2", "3.54%", "₩563", "₩28,699", "7", "₩4,099"],
            ["AIMBTI-fear-v2", "2.73%", "₩634", "₩31,710", "1", "₩31,710"],
            ["합계 (활성 3개)", "3.04%", "₩560", "₩88,564", "15", "₩5,904"],
        ],
    )

    add_subheading(doc, "■ 핵심 인사이트")
    add_bullets(doc, [
        "Simple 소재: CPL ₩4,022 · CTR 2.96% — 현재 최우수, 예산 집중 배분",
        "Social 소재: CPL ₩4,099 · CTR 3.54% — Simple과 거의 동률, 안정적",
        "Fear 소재: Lead 1건 · CPL ₩31,710 — 감정 소구 한계 검증 후 조기 정지 판단",
        "평균 CTR 3.04% — Meta 업계 평균(1.91%) 대비 59% 우수 (WordStream 2024 기준)",
        "Meta Pixel Lead는 세션당 1회만 발사하도록 설계 — Ads Manager 수치와 일치",
    ])

    page_break(doc)

    # ─────────────────────────────────────────────
    # 6. 로드맵 & Contact
    # ─────────────────────────────────────────────
    add_heading(doc, "05. 로드맵", size=22)

    add_subheading(doc, "■ 단기 (2026 Q2)")
    add_bullets(doc, [
        "Meta 3-캠페인 A/B/C 소재 고도화 및 Simple 소재 예산 확대",
        "랜딩페이지 v5 제작 — v4 데이터 기반 CTA 재배치",
        "결과 페이지 CTA 전환율 개선 (이북 언락 플로우 UX 개선)",
    ])

    add_subheading(doc, "■ 중기 (2026 Q3~Q4)")
    add_bullets(doc, [
        "유형별 부트캠프·이북 콘텐츠 상품 확장",
        "B2B 기업 진단 프로덕트 — 팀 단위 AI 역량 분석 대시보드",
        "캐릭터 IP 기반 오가닉 콘텐츠 시리즈 정규화 (주간 발행)",
    ])

    add_subheading(doc, "■ 장기")
    add_bullets(doc, [
        "영문·일문 버전 출시 → 글로벌 AI 교육 시장 진입",
        "진단 데이터 기반 AI 역량 벤치마크 리포트 발간",
        "기업 HR·L&D 부서 대상 SaaS 대시보드 별도 상품화",
    ])

    add_heading(doc, "Contact", size=18)
    add_para(doc, "서비스: https://mcodegc.com", size=11)
    add_para(doc, "IR/PR 문의: 내부 담당자 경유", size=11, color=GRAY)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✓ 생성 완료: {OUT}")
    print(f"  크기: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
