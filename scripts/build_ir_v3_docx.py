"""AI-MBTI IR/PR v3 — 비전문가용 쉬운 말 버전 docx.

실행: python3 scripts/build_ir_v3_docx.py
출력: docs/IR_AI-MBTI_v3.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
OUT = ROOT / "docs" / "IR_AI-MBTI_v3.docx"

NAVY = RGBColor(0x0A, 0x1B, 0x3D)
ACCENT = RGBColor(0x1E, 0x6B, 0xFF)
TEXT = RGBColor(0x1A, 0x1F, 0x2E)
MUTED = RGBColor(0x6B, 0x72, 0x80)
SUCCESS = RGBColor(0x0E, 0xA3, 0x6B)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def section_header(doc, label: str, title: str, subtitle: str | None = None):
    # 컬러바
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("▬▬")
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT
    run.bold = True

    # 섹션 라벨
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = ACCENT

    # 섹션 타이틀
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.bold = True
    run.font.color.rgb = NAVY

    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT


def add_h3(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY


def add_para(doc, text: str, size: int = 10, bold: bool = False, color: RGBColor | None = None,
             align: int | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bullets(doc, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0] if p.runs else p.add_run()
        run.text = it
        run.font.size = Pt(10)


def add_image(doc, rel_path: str, width_cm: float = 15.0, caption: str | None = None) -> None:
    img = PUBLIC / rel_path
    if not img.exists():
        add_para(doc, f"[이미지 없음: {rel_path}]", color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED


def add_image_row(doc, entries: list[tuple[str, str]], row_width_cm: float = 16.0) -> None:
    valid = [(PUBLIC / p, c) for p, c in entries if (PUBLIC / p).exists()]
    if not valid:
        return
    n = len(valid)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (img_path, caption) in enumerate(valid):
        cell_img = tbl.cell(0, i)
        cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell_img.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(str(img_path), width=Cm(row_width_cm / n - 0.3))
        except Exception as e:
            para.add_run(f"[err {img_path.name}]").font.size = Pt(8)
        cap_p = tbl.cell(1, i).paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap_p.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = MUTED


def add_kpi_row(doc, cards: list[tuple[str, str, str]]) -> None:
    """큰 숫자 KPI 카드 3~4개 가로 배치."""
    n = len(cards)
    tbl = doc.add_table(rows=3, cols=n)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (label, big, sub) in enumerate(cards):
        # row 0: label
        c0 = tbl.cell(0, i)
        set_cell_bg(c0, "F3F6FC")
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED
        # row 1: big
        c1 = tbl.cell(1, i)
        set_cell_bg(c1, "F3F6FC")
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(big)
        r.font.size = Pt(26)
        r.bold = True
        r.font.color.rgb = NAVY
        # row 2: sub
        c2 = tbl.cell(2, i)
        set_cell_bg(c2, "F3F6FC")
        p = c2.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sub)
        r.font.size = Pt(9)
        r.font.color.rgb = SUCCESS


def add_table(doc, headers: list[str], rows: list[list[str]], highlight_rows: list[int] | None = None) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "0A1B3D")
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            if highlight_rows and ri in highlight_rows:
                r.bold = True
                r.font.color.rgb = NAVY
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "F3F6FC")


def add_pullquote(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"“{text}”")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY


def page_break(doc) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(10)

    # ══════════ 표지 ══════════
    for _ in range(3):
        doc.add_paragraph()
    logo = PUBLIC / "Adv" / "로고 및 참조" / "로고_transparent.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Cm(6))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI ERA CAREER DIAGNOSIS")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-MBTI")
    r.bold = True
    r.font.size = Pt(54)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 시대, 당신의 생존력을 진단합니다")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("20문항으로 AI 활용 유형 진단 · 데이터 기반 광고 운영")
    r.font.size = Pt(13)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    add_kpi_row(doc, [
        ("광고비", "₩88,564", "Meta 활성 3개 · 2일간"),
        ("잠재고객", "16명", "중복 제거 기준"),
        ("1명 확보 비용", "₩4,022", "업계 대비 1/6 수준"),
    ])

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IR / PR Deck v2 · 2026.04 · mcodegc.com")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    page_break(doc)

    # ══════════ 01. Executive Summary ══════════
    section_header(doc, "01  EXECUTIVE SUMMARY", "한 페이지 요약",
                   "AI-MBTI는 직장인의 AI 활용 역량을 20문항으로 진단하고, 결과 기반 학습 경로를 제안하는 서비스입니다. "
                   "2026년 4월 기준 누적 4,339명이 방문했으며, 소재 A/B 테스트로 최우수 소재(Social)의 CPL을 ₩3,980까지 최적화한 상태로 운영되고 있습니다.")

    add_h3(doc, "핵심 경쟁력")
    add_bullets(doc, [
        "진단 엔진 — 5차원 점수제 → 16유형 분류 (고유 로직)",
        "데이터 인프라 — Meta Ads + Google Ads + Supabase 단일 대시보드 통합",
        "소재 최적화 역량 — A/B/C 테스트로 최우수 Social 소재 CPL ₩3,980 확보 (Meta 업계 평균 약 ₩25,000 대비 84% 저렴)",
        "IP 자산 — 16유형 캐릭터 디자인 보유, 2차 저작물·굿즈 확장 가능",
    ])

    add_h3(doc, "비즈니스 기회")
    add_bullets(doc, [
        "국내 AI 교육 시장 급성장 — 직장인 AI 리스킬링 수요 폭증",
        "B2B 확장 — 기업 HR·L&D 대상 팀 역량 진단 상품화",
        "콘텐츠 자산 재활용 — 캐릭터 IP 기반 굿즈·주간 카드뉴스·커뮤니티",
    ])

    doc.add_paragraph()
    add_pullquote(doc, "A/B 테스트로 최우수 소재 CPL ₩3,980을 찾았습니다.")

    page_break(doc)

    # ══════════ 02. Problem & Solution ══════════
    section_header(doc, "02  PROBLEM & SOLUTION", "문제와 해결",
                   "AI 시대 직장인은 막연한 불안을 느끼지만 구체적으로 무엇을 배워야 하는지 모릅니다. "
                   "AI-MBTI는 이 불안을 구조화된 진단으로 해소합니다.")

    add_para(doc, "PROBLEM", size=11, bold=True, color=ACCENT)
    add_h3(doc, "AI가 일상·업무에 깊숙이 들어왔지만 개인의 대응은 제각각")
    add_bullets(doc, [
        "“나는 AI를 어느 정도 쓸 수 있는 사람인가?”에 대한 객관적 답 부재",
        "수많은 강의·부트캠프 중 무엇이 내게 맞는지 판단 근거 없음",
        "막연한 불안 → 소비 전환으로 이어지지 않는 미결정 상태 장기화",
    ])

    doc.add_paragraph()
    add_para(doc, "SOLUTION", size=11, bold=True, color=ACCENT)
    add_h3(doc, "20문항 · 3분 진단 → 16유형 중 나의 위치 판정 → 맞춤 학습 경로")

    add_image_row(doc, [
        ("characters/AI 팀장.png", "AI 활용형"),
        ("characters/LLM 설계자.png", "LLM 전문형"),
        ("characters/DE 전략가.png", "데이터 전략형"),
        ("characters/DA 예술가.png", "창작 활용형"),
    ], row_width_cm=16)

    doc.add_paragraph()
    add_para(doc,
             "각 유형은 고유 캐릭터 · 강점 · 추천 학습 트랙을 가지며, 결과 페이지에서 "
             "전자책·부트캠프·오픈챗으로 자연스럽게 연결됩니다.",
             size=10)

    page_break(doc)

    # ══════════ 03. Traction ══════════
    section_header(doc, "01  성과", "숫자로 보는 현재",
                   "2026년 4월 22일 ~ 4월 23일, 페이스북·인스타 광고 3종 집행 결과입니다.")

    add_kpi_row(doc, [
        ("광고비", "₩88,813", "2일간 집행"),
        ("관심 고객", "17명", "중복 제거"),
        ("1명 모시는 비용", "₩4,030", "최우수 광고"),
        ("광고 클릭률", "3.03%", "업계의 1.6배"),
    ])

    doc.add_paragraph()
    add_h3(doc, "캠페인 3종 성적표 (여성 이미지 × 첫페이지 3종, 4월 22~23일)")
    add_table(doc,
              ["캠페인 (여성 × 첫페이지)", "관심 고객", "1명 모시는 비용", "한 줄 평가"],
              [
                  ["여성 × 심플 첫페이지 (1등)", "7명", "₩4,030", "정보 중심 첫페이지 · 가장 저렴"],
                  ["여성 × 사회적 첫페이지", "7명", "₩4,121", "공감형 첫페이지 · 클릭 가장 많음"],
                  ["여성 × 공포 첫페이지 (실험)", "3명", "₩10,585", "공포 소구, 효과 애매해서 중단 예정"],
                  ["합계 (3종)", "17명", "평균 ₩5,224", "심플·사회적 첫페이지 중심 운영"],
              ],
              highlight_rows=[0])

    doc.add_paragraph()
    add_h3(doc, "우리 광고 vs 업계 평균")
    add_table(doc,
              ["지표", "업계 평균", "우리", "차이"],
              [
                  ["광고 보고 누르는 비율", "1.91%", "3.03%", "▲ 1.6배"],
                  ["한 번 누를 때 광고비", "₩1,100", "₩562", "▼ 절반"],
                  ["1명 모시는 비용", "약 ₩25,000", "₩4,030", "▼ 1/6"],
              ])
    add_para(doc,
             "출처: WordStream 2024 Meta 광고 벤치마크 · 국내 교육·컨설팅 업종 기준. 기간: 4월 22일~23일.",
             size=9, color=MUTED)

    doc.add_paragraph()
    add_pullquote(doc, "업계 평균의 1/6 비용으로 관심 고객을 모으고 있습니다.")

    page_break(doc)

    # ══════════ 관심 고객 + 규모별 광고비 페이지 ══════════
    add_h3(doc, "관심 고객 17명 — 어디로 갔나? (4월 22~23일)")
    add_table(doc,
              ["캠페인 (첫페이지)", "총 관심 고객", "전자책으로 이동", "단톡방으로 이동"],
              [
                  ["Simple", "7명", "7명", "0명"],
                  ["Social", "7명", "6명", "1명"],
                  ["Fear", "3명", "2명", "1명"],
                  ["합계", "17명", "15명", "2명"],
              ],
              highlight_rows=[3])
    add_para(doc,
             "※ Meta 광고 관리자 기준 관심 고객 수 (같은 사람이 여러 번 눌러도 1명으로 집계). "
             "전자책 이동 = 미리보기 끝까지 본 사람 or 다운로드 버튼 누른 사람. "
             "단톡방 이동 = '단톡방 입장' 버튼 누른 사람.",
             size=9, color=MUTED)
    add_bullets(doc, [
        "Simple (7명): 전원 전자책 · 정보 원하는 사람 모임",
        "Social (7명): 전자책 6 + 단톡방 1 · 가장 다양한 반응",
        "Fear (3명): 전자책 2 + 단톡방 1 · 수는 적지만 단톡방까지",
        "전체: 17명 중 15명(88%)이 전자책 → 전자책이 핵심 유도 상품",
    ])

    doc.add_paragraph()
    add_h3(doc, "규모별 광고비 환산 — 고객 N명 모으려면 얼마?")
    add_table(doc,
              ["규모", "현재 CPL 기준 (₩4,030)", "실질 가격 (150%)", "해석"],
              [
                  ["10명", "₩40,300", "₩60,450", "테스트 규모"],
                  ["100명", "₩403,000", "₩604,500", "월 소규모 운영"],
                  ["1,000명", "₩4,030,000", "₩6,045,000", "월 중형 캠페인"],
                  ["10,000명", "₩40,300,000", "₩60,450,000", "연간 대규모 확장"],
              ],
              highlight_rows=[3])
    add_para(doc,
             "※ 150% 할증 근거: 광고 규모를 10~10,000배 확대하면 최적 타겟 고갈로 CPL 상승. "
             "업계 표준상 초기 CPL의 130~170%로 수렴 → 중간값 150%를 보수적으로 적용.",
             size=9, color=MUTED)

    page_break(doc)

    # ══════════ Meta 캠페인 비교 페이지 ══════════
    add_h3(doc, "Meta 광고 캠페인별 전체 퍼널 비교 (4월 22~23일)")
    add_para(doc,
             "페이스북·인스타 광고로 들어온 사람들만 추려서 (utm_source=meta) 캠페인별로 "
             "어디에서 얼마나 이탈하는지 비교한 결과입니다.", size=10)
    add_table(doc,
              ["캠페인", "랜딩 조회", "CTA 클릭", "테스트 시작", "결과 확인", "Lead", "CTA 클릭률 ★", "Lead 전환율 ★"],
              [
                  ["Simple", "92", "51", "44", "37", "7", "55.4% 🏆 1위", "7.6%"],
                  ["Social", "73", "36", "33", "32", "7", "49.3%", "9.6% 🏆 1위"],
                  ["Fear", "84", "45", "38", "29", "3", "53.6% (2위)", "3.6% (꼴찌)"],
                  ["합계·평균", "249", "132", "115", "98", "17", "53.0%", "6.8%"],
              ],
              highlight_rows=[3])
    add_para(doc,
             "★ 두 개의 '1위'가 서로 다른 캠페인: CTA 클릭률 1위 = Simple (55.4%) / Lead 전환율 1위 = Social (9.6%). "
             "버튼 많이 누르게 하는 광고(Simple)와 끝까지 전환되는 광고(Social)가 따로임.",
             size=9, color=MUTED)

    doc.add_paragraph()
    add_h3(doc, "결론 — 3개 광고의 역할이 명확히 다름")
    add_bullets(doc, [
        "🏆 Simple = 트래픽 수확기: CTA 클릭률 55.4% 1위. 넓은 그물로 많이 끌어오는 메인 소재",
        "🏆 Social = 질적 전환기: Lead 전환율 9.6% 1위. 관심 없는 사람은 걸러지고 진짜 관심 있는 사람만 끝까지 감",
        "⚠ Fear = 호기심 낚시 (중단 예정): Lead 전환율 3.6% 꼴찌. 클릭은 유도되나 실제 전환 안 됨",
    ])

    page_break(doc)

    # ══════════ Meta 캠페인 상세 비교 (데이터 분석 홍보 포함) ══════════
    add_h3(doc, "Meta 광고 캠페인 비교 — AI-MBTI vs 데이터 분석 홍보 (같은 계정·같은 기간)")
    add_para(doc,
             "같은 메타코드M 광고 계정에서 돌리는 '데이터 분석 홍보' 캠페인과 AI-MBTI 3개를 "
             "동일 기간(4월 22~23일)·동일 조건(OUTCOME_LEADS)으로 비교.", size=10)
    add_table(doc,
              ["지표", "Simple (1위)", "Social", "Fear (실험)", "데이터 분석 홍보"],
              [
                  ["광고 목표", "잠재고객 확보", "잠재고객 확보", "잠재고객 확보", "잠재고객 확보"],
                  ["광고비", "₩29,891", "₩30,260", "₩33,593", "₩95,408"],
                  ["도달 (실사용자)", "1,526명", "1,242명", "1,495명", "2,627명"],
                  ["노출 수", "2,052회", "1,528회", "1,929회", "3,067회"],
                  ["클릭률 (CTR)", "3.02%", "3.34%", "2.80%", "2.22%"],
                  ["클릭당 비용 (CPC)", "₩482", "₩593", "₩622", "₩1,403"],
                  ["Lead", "9명", "8명", "5명", "11명"],
                  ["1명 모시는 비용 (CPL)", "₩3,321 🏆", "₩3,783", "₩6,719", "₩8,673"],
              ],
              highlight_rows=[7])
    add_para(doc,
             "※ Meta Graph API v21.0 직접 조회 (2026-04-22~04-23). 동일 계정·동일 광고 목표로 공정 비교 가능.",
             size=9, color=MUTED)

    doc.add_paragraph()
    add_h3(doc, "보이는 인사이트 — AI-MBTI가 압도적 우위")
    add_bullets(doc, [
        "CPL 비교: AI-MBTI Simple ₩3,321 vs 데이터 분석 홍보 ₩8,673 — 2.6배 저렴",
        "CPC 비교: Simple ₩482 vs 데이터 분석 홍보 ₩1,403 — 2.9배 저렴",
        "CTR 비교: AI-MBTI 평균 3.05% vs 데이터 분석 홍보 2.22% — +37% 우수",
        "효율 역전: 데이터 분석 홍보가 광고비 3배 더 쓰고도 Lead 11명 (AI-MBTI Simple 9명과 거의 동일)",
    ])

    page_break(doc)

    # ══════════ 04. Product ══════════
    section_header(doc, "04  PRODUCT", "제품 구조",
                   "진단 → 결과 → 학습으로 이어지는 3단계 플로우.")

    # 3-step flow
    tbl = doc.add_table(rows=3, cols=3)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    steps_data = [
        ("STEP 1", "진단", "20문항 3분"),
        ("STEP 2", "결과", "16유형 중 매칭"),
        ("STEP 3", "학습 연결", "부트캠프·이북·오픈챗"),
    ]
    for i, (top, mid, bot) in enumerate(steps_data):
        for ri, (text, size, bold, color) in enumerate([
            (top, 9, False, MUTED),
            (mid, 16, True, NAVY),
            (bot, 9, False, MUTED),
        ]):
            cell = tbl.cell(ri, i)
            set_cell_bg(cell, "F3F6FC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = bold
            r.font.color.rgb = color

    doc.add_paragraph()
    add_h3(doc, "랜딩 페이지 — 3버전 동시 운영")
    add_image_row(doc, [
        ("landing/v1/full.png", "v1"),
        ("landing/v3/full.png", "v3"),
        ("landing/v4/full.png", "v4 (모바일)"),
    ], row_width_cm=14)

    page_break(doc)

    # ══════════ 05. Marketing ══════════
    section_header(doc, "05  MARKETING", "마케팅 실행",
                   "Meta Ads 중심 퍼포먼스 마케팅 — 소재·오디언스·랜딩 3축 동시 실험.")

    add_h3(doc, "소재 포트폴리오")
    add_image_row(doc, [
        ("Adv/여자/사무실 여자 완성.png", "Simple · 여성 타겟"),
        ("Adv/남자/사무실 남자 완성.png", "Simple · 남성 타겟"),
        ("screen/fear/fear.png", "Fear (실험)"),
    ], row_width_cm=17)

    doc.add_paragraph()
    add_bullets(doc, [
        "성별 타겟팅 이원화 — 남성·여성 소재 독립 제작으로 오디언스별 전환율 측정",
        "9:16 · 1:1 · 1.91:1 3개 규격 동시 운영 — 전 지면 커버",
        "감정 vs 실용 소구 비교 — Fear 캠페인으로 한계 데이터 확보",
        "주 1회 리뷰 사이클 — 하위 성과 정리 · 상위 예산 확대",
    ])

    doc.add_paragraph()
    add_h3(doc, "오가닉 채널 — 캐릭터 IP 활용 카드뉴스")
    add_image_row(doc, [
        ("zunza/AISERVICE/1.png", "AI 서비스"),
        ("zunza/DA/1.png", "DA 시리즈"),
        ("zunza/AILLM/1.png", "LLM 시리즈"),
    ], row_width_cm=16)

    page_break(doc)

    # ══════════ 06. Data Infrastructure ══════════
    section_header(doc, "06  DATA INFRASTRUCTURE", "데이터 인프라",
                   "의사결정의 속도가 경쟁력입니다. Meta · Google · Supabase를 단일 뷰로 통합한 "
                   "자체 대시보드 v3를 운영 중입니다.")

    add_h3(doc, "통합 대시보드 — 종합 뷰")
    add_image(doc, "screen/dashboard3.png", width_cm=16,
              caption="퍼널 · 광고 성과 · 진단 엔진 통합 — 위험/주의/양호 라벨 실시간")

    page_break(doc)

    add_h3(doc, "퍼널 분석 뷰 — 업계 벤치마크 자동 비교")
    add_image(doc, "screen/dashboard3-funnel.png", width_cm=16,
              caption="Unbounce 2024 · WordStream 2025 · LeadQuizzes 벤치마크 내장")

    doc.add_paragraph()
    add_h3(doc, "핵심 전환율 vs 업계 평균")
    add_table(doc,
              ["구간", "당사", "업계 평균", "차이"],
              [
                  ["방문 → CTA 클릭", "27.6%", "40.0%", "▼ -31% (개선 여지)"],
                  ["CTA → 테스트 시작", "91.0%", "90.0%", "▲ +1%"],
                  ["테스트 시작 → 완료", "95.7%", "70.0%", "▲ +37%"],
                  ["테스트 완료 → 결과 확인", "93.9%", "95.0%", "▼ -1%"],
                  ["결과 확인 → 2차 전환", "39.8%", "8.0%", "▲ +398% ★"],
              ],
              highlight_rows=[4])

    doc.add_paragraph()
    add_para(doc,
             "★ 결과 확인 → 2차 전환 39.8%는 업계 평균 8% 대비 5배 — "
             "진단 UX가 자연스럽게 학습 상품 구매로 이어지고 있음을 입증합니다.",
             size=10)

    page_break(doc)

    add_h3(doc, "Meta Ads 뷰 — 8개 KPI 통합")
    add_image(doc, "screen/dashboard3-meta.png", width_cm=16,
              caption="지출 · 노출 · 클릭 · CTR · CPC · CPM · 전환 · CVR")

    doc.add_paragraph()
    add_h3(doc, "A/B 테스트 뷰 — 변형별 즉시 비교")
    add_image(doc, "screen/dashboard3-abtest.png", width_cm=16,
              caption="동일 기간 변형별 독립 집계 → 예산 재배분 근거")

    page_break(doc)

    # ══════════ 07. Roadmap ══════════
    section_header(doc, "07  ROADMAP", "로드맵",
                   "단기 수익화 · 중기 B2B 확장 · 장기 글로벌 진출.")

    add_table(doc,
              ["시기", "핵심 목표", "주요 과제"],
              [
                  ["2026 Q2", "CPL 추가 단축 · 전환율 개선",
                   "Simple 소재 예산 확대 · 랜딩 v5 · 결과 페이지 CTA 개선"],
                  ["2026 Q3-Q4", "상품 다양화 · B2B 진입",
                   "유형별 부트캠프 · 기업 팀 진단 SaaS · 주간 콘텐츠 정규화"],
                  ["2027+", "글로벌 · 벤치마크 리포트",
                   "영·일문 버전 · AI 역량 연간 리포트 · HR/L&D 대상 엔터프라이즈"],
              ])

    doc.add_paragraph()
    doc.add_paragraph()
    section_header(doc, "", "연락처")
    add_para(doc, "서비스: https://mcodegc.com", size=11)
    add_para(doc, "IR / PR 문의: information@mcode.co.kr", size=11)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✓ 생성 완료: {OUT}")
    print(f"  크기: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
